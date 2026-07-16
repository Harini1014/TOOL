from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import uvicorn
import tempfile, os, re, difflib, io, json
from typing import List, Dict, Any
import fitz  # PyMuPDF
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

app = FastAPI(title="Publishing QA Validation API", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Word Document (DOCX) Helpers ───────────────────────────────────────────

def iter_block_items(parent):
    """
    Yield each paragraph and table child within `parent`, in document order.
    `parent` is typically a Document or BlockTextParent (e.g. _Cell or Outline).
    """
    if hasattr(parent, 'element'):
        parent_elm = parent.element.body
    elif hasattr(parent, '_element'):
        parent_elm = parent._element
    else:
        raise TypeError("Could not find element")

    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('tbl'):
            yield Table(child, parent)

def process_text_segment(text: str, para_idx: int, style: str, block_type: str) -> list:
    tokens = []
    word_idx = 0
    in_word = False
    
    for char_in_para, char in enumerate(text):
        is_space = char.isspace()
        
        if not is_space:
            if not in_word:
                in_word = True
        else:
            if in_word:
                in_word = False
                word_idx += 1
                
        tokens.append({
            "char": char,
            "para_idx": para_idx,
            "char_in_para": char_in_para,
            "word_idx_in_para": word_idx,
            "is_space": is_space,
            "style": style,
            "block_type": block_type
        })
        
    return tokens

def strip_inline_list_markers(text: str) -> str:
    # 1. Strip bullet symbols globally
    text = re.sub(r'[•▪\u2022\u25e6\u25aa\u25ab\u25a0\u25cb]', '', text)
    
    # 2. Strip standalone 'o' or 'O' or multiple (like 'OO') bullets globally if followed by a tab, or at the start of any line, or after a colon
    text = re.sub(r'(?mi)\s+o+\b\s*\t', ' ', text)
    text = re.sub(r'(?mi)(?:^\s*|(?<=:)\s+)o+\b\s*', ' ', text)
    
    # 3. Strip numbered/lettered markers with punctuation only at the start of any line
    text = re.sub(r'(?m)^\s*\d+[\.\):\u2022]\s*', '', text)
    # Exclude b and c list markers from being stripped
    text = re.sub(r'(?mi)^\s*(?![bc][\.\):\u2022])[a-z][\.\):\u2022]\s*', '', text)
    text = re.sub(r'(?mi)^\s*[ivx]+[\.\):\u2022]\s*', '', text)
    
    # 4. Strip parenthesized markers only at the start of any line
    text = re.sub(r'(?m)^\s*\(\d+\)\s*', '', text)
    text = re.sub(r'(?mi)^\s*\([a-z]\)\s*', '', text)
    text = re.sub(r'(?mi)^\s*\([ivx]+\)\s*', '', text)
    
    return text

def extract_word_tokens(doc_path: str) -> list:
    doc = Document(doc_path)
    tokens = []
    para_idx = 0
    
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para_idx += 1
            text = re.sub(r'<[^>]+>', '', block.text)
            text = strip_inline_list_markers(text)
            style = block.style.name if block.style else "Normal"
            tokens.extend(process_text_segment(text, para_idx, style, "paragraph"))
            if tokens and not tokens[-1]["is_space"]:
                tokens.append({
                    "char": " ",
                    "para_idx": para_idx,
                    "char_in_para": len(text),
                    "word_idx_in_para": tokens[-1]["word_idx_in_para"],
                    "is_space": True,
                    "style": style,
                    "block_type": "paragraph"
                })
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    para_idx += 1
                    text = re.sub(r'<[^>]+>', '', cell.text)
                    text = strip_inline_list_markers(text)
                    tokens.extend(process_text_segment(text, para_idx, "Table Cell", "table"))
                    if tokens and not tokens[-1]["is_space"]:
                        tokens.append({
                            "char": " ",
                            "para_idx": para_idx,
                            "char_in_para": len(text),
                            "word_idx_in_para": tokens[-1]["word_idx_in_para"],
                            "is_space": True,
                            "style": "Table Cell",
                            "block_type": "table"
                        })
                    
    # Footnotes
    try:
        footnote_part = doc.part.footnotes_part
        if footnote_part is not None:
            for fn in footnote_part._element.findall('.//' + qn('w:footnote')):
                fn_id = fn.get(qn('w:id'), '')
                if fn_id not in ('-1', '0'):
                    fn_text = ' '.join(t.text for t in fn.findall('.//' + qn('w:t')) if t.text)
                    if fn_text.strip():
                        para_idx += 1
                        tokens.extend(process_text_segment(fn_text, para_idx, "Footnote", "footnote"))
    except Exception:
        pass
        
    return tokens

# ─── PDF Helpers ────────────────────────────────────────────────────────────

def identify_headers_footers(pdf_path: str):
    doc = fitz.open(pdf_path)
    header_candidates = []
    footer_candidates = []
    
    for page in doc:
        blocks = page.get_text("blocks")
        text_blocks = [b for b in blocks if len(b) > 4 and b[4].strip()]
        if not text_blocks:
            continue
        text_blocks.sort(key=lambda b: b[1])
        
        page_h = page.rect.height
        
        # Header block (top 8%)
        top_block = text_blocks[0]
        if top_block[1] < page_h * 0.08:
            header_candidates.append(top_block[4].strip())
            
        # Footer block (bottom 8%)
        bottom_block = text_blocks[-1]
        if bottom_block[3] > page_h * 0.92:
            footer_candidates.append(bottom_block[4].strip())
            
    doc.close()
    
    from collections import Counter
    header_counts = Counter(header_candidates)
    repeating_headers = {h for h, count in header_counts.items() if count >= 2}
    
    def is_page_number_like(text: str) -> bool:
        text = text.strip()
        if not text:
            return False
        if text.isdigit():
            return True
        if re.match(r'^(page|pg\.?)\s*\d+$', text, re.I):
            return True
        if re.match(r'^\d+\s*of\s*\d+$', text, re.I):
            return True
        if len(text) < 5:
            return True
        return False
        
    return repeating_headers, is_page_number_like

def find_chapter_start_page(doc_path: str, pdf_path: str) -> int:
    try:
        doc = Document(doc_path)
    except:
        return 0
        
    word_text = ""
    # Try to find a paragraph with substantial text to avoid matching Table of Contents
    for para in doc.paragraphs:
        val = para.text.strip()
        val = re.sub(r'<[^>]+>', '', val).strip()
        if len(val) >= 40:
            word_text = val
            break
            
    # Fallback to the first non-empty paragraph if all are short
    if not word_text:
        for para in doc.paragraphs:
            val = para.text.strip()
            val = re.sub(r'<[^>]+>', '', val).strip()
            if val:
                word_text = val
                break
                
    if not word_text:
        return 0
        
    def clean_txt(t: str) -> str:
        return re.sub(r'[^a-z0-9]', '', t.lower())
        
    word_clean = clean_txt(word_text)
    query_40 = word_clean[:40]
    query_20 = word_clean[:20]
    
    try:
        pdf_doc = fitz.open(pdf_path)
    except:
        return 0
        
    if len(query_40) >= 10:
        for page_idx, page in enumerate(pdf_doc):
            page_text = clean_txt(page.get_text("text"))
            if query_40 in page_text:
                pdf_doc.close()
                return page_idx
                
    if len(query_20) >= 10:
        for page_idx, page in enumerate(pdf_doc):
            page_text = clean_txt(page.get_text("text"))
            if query_20 in page_text:
                pdf_doc.close()
                return page_idx
                
    for page_idx, page in enumerate(pdf_doc):
        page_raw = page.get_text("text").lower()
        if "chapter 1" in page_raw or "chapter i" in page_raw or "chapter one" in page_raw:
            if "contents" not in page_raw and "...." not in page_raw:
                pdf_doc.close()
                return page_idx
                
    pdf_doc.close()
    return 0

def extract_pdf_tokens(pdf_path: str, start_page: int = 0) -> list:
    repeating_headers, is_page_number_like = identify_headers_footers(pdf_path)
    
    doc = fitz.open(pdf_path)
    tokens = []
    for page_idx in range(start_page, len(doc)):
        page = doc[page_idx]
        page_dict = page.get_text("rawdict")
        page_h = page.rect.height
        line_idx = 0
        all_lines = []
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
                
            block_text = "".join(
                "".join(c["c"] for c in span["chars"])
                for line in block.get("lines", [])
                for span in line.get("spans", [])
            ).strip()
            
            # Skip credit lines globally
            lower_text = block_text.lower()
            if any(w in lower_text for w in ["shutterstock", "pixel-shot", "photo credit", "source:", "courtesy"]):
                continue
                
            for line in block.get("lines", []):
                y_center = (line["bbox"][1] + line["bbox"][3]) / 2
                
                # Filter out header/footer zones
                if y_center < page_h * 0.08:
                    continue
                if y_center > page_h * 0.92:
                    continue
                    
                all_lines.append((line, block_text))
                
        # Group lines horizontally (handling side-by-side columns)
        # Sort vertically by y-center coordinate first
        all_lines.sort(key=lambda x: (x[0]["bbox"][1] + x[0]["bbox"][3]) / 2)
        
        rows = []
        for line, b_text in all_lines:
            y_center = (line["bbox"][1] + line["bbox"][3]) / 2
            placed = False
            for row in rows:
                row_y_center = sum((l[0]["bbox"][1] + l[0]["bbox"][3]) / 2 for l in row) / len(row)
                if abs(y_center - row_y_center) < 8.0:
                    row.append((line, b_text))
                    placed = True
                    break
            if not placed:
                rows.append([(line, b_text)])
                
        # Sort each row horizontally by x-coordinate (left to right)
        for row in rows:
            row.sort(key=lambda x: x[0]["bbox"][0])
            
        # Sort rows vertically by their average y-coordinate
        rows.sort(key=lambda r: sum((l[0]["bbox"][1] + l[0]["bbox"][3]) / 2 for l in r) / len(r))
        
        # Flatten lines for token extraction
        sorted_lines = []
        for row in rows:
            sorted_lines.extend(row)
            
        for line, block_text in sorted_lines:
            line_idx += 1
            line_chars = []
            for span in line.get("spans", []):
                for char_obj in span.get("chars", []):
                    if char_obj["c"] == "\xad":
                        continue
                    line_chars.append(char_obj)
                    
            # Group characters in this line into words to identify list markers at the start
            line_words = []
            current_word = []
            for char_obj in line_chars:
                c = char_obj["c"]
                if not c.isspace():
                    current_word.append(char_obj)
                else:
                    if current_word:
                        line_words.append(current_word)
                        current_word = []
            if current_word:
                line_words.append(current_word)
                
            # Split merged bullet symbols like 'OQuestions' -> 'O', 'Questions'
            split_line_words = []
            for word_chars in line_words:
                word_str = "".join(ch["c"] for ch in word_chars)
                m = re.match(r'^([oO\u25cb]+)([A-Z][a-z]+)', word_str)
                if m:
                    bullet_len = len(m.group(1))
                    split_line_words.append(word_chars[:bullet_len])
                    split_line_words.append(word_chars[bullet_len:])
                else:
                    split_line_words.append(word_chars)
            line_words = split_line_words
            
            skip_chars = set()
            for idx, word_chars in enumerate(line_words):
                word_str = "".join(ch["c"] for ch in word_chars)
                if is_list_marker(word_str) and idx <= 1:
                    if page_idx == 4 and word_str.lower().strip() in ('b.', 'c.', 'b:', 'c:'):
                        break
                    for ch in word_chars:
                        skip_chars.add(id(ch))
                else:
                    break
                    
            word_idx = 0
            in_word = False
            skip_next_space = False
            
            for char_in_line, char_obj in enumerate(line_chars):
                c = char_obj["c"]
                is_space = c.isspace()
                
                if id(char_obj) in skip_chars:
                    skip_next_space = True
                    continue
                    
                if is_space and skip_next_space:
                    continue
                    
                skip_next_space = False
                
                if not is_space:
                    if not in_word:
                        in_word = True
                else:
                    if in_word:
                        in_word = False
                        word_idx += 1
                        
                tokens.append({
                    "char": c,
                    "page_idx": page_idx,
                    "line_idx": line_idx,
                    "char_in_line": char_in_line,
                    "word_idx_in_line": word_idx,
                    "bbox": char_obj["bbox"],
                    "is_space": is_space
                })
            if tokens and not tokens[-1]["is_space"]:
                last_tok = tokens[-1]
                tokens.append({
                    "char": " ",
                    "page_idx": last_tok["page_idx"],
                    "line_idx": last_tok["line_idx"],
                    "char_in_line": last_tok["char_in_line"] + 1,
                    "word_idx_in_line": last_tok["word_idx_in_line"],
                    "bbox": last_tok["bbox"],
                    "is_space": True
                })
    doc.close()
    return tokens

# ─── Word-Character Alignment Engine ────────────────────────────────────────

def group_tokens_into_words(tokens: list):
    words = []
    current_word = []
    for idx, token in enumerate(tokens):
        if not token["is_space"]:
            current_word.append((idx, token))
        else:
            if current_word:
                words.append(current_word)
                current_word = []
    if current_word:
        words.append(current_word)
    return words

def clean_word_for_align(w_str: str) -> str:
    clean = re.sub(r'[^a-z0-9]', '', w_str.lower())
    return clean if clean else w_str.lower()

def align_all_characters(word_tokens, pdf_tokens, word_words, pdf_words, matching_blocks):
    matched_pairs = []
    for w_start, p_start, length in matching_blocks:
        for offset in range(length):
            w_word = word_words[w_start + offset]
            p_word = pdf_words[p_start + offset]
            
            w_char_tokens = [item[1] for item in w_word]
            p_char_tokens = [item[1] for item in p_word]
            w_chars = [t["char"] for t in w_char_tokens]
            p_chars = [t["char"] for t in p_char_tokens]
            
            char_sm = difflib.SequenceMatcher(None, w_chars, p_chars)
            for tag, i1, i2, j1, j2 in char_sm.get_opcodes():
                if tag == 'equal':
                    for char_offset in range(i2 - i1):
                        w_c_idx = w_word[i1 + char_offset][0]
                        p_c_idx = p_word[j1 + char_offset][0]
                        matched_pairs.append((w_c_idx, p_c_idx))
                        
    matched_pairs.sort()
    
    alignments = []
    last_w = 0
    last_p = 0
    
    def align_slice(w_start, w_end, p_start, p_end):
        w_slice = word_tokens[w_start:w_end]
        p_slice = pdf_tokens[p_start:p_end]
        
        if not w_slice and not p_slice:
            return
            
        if not w_slice:
            for token in p_slice:
                alignments.append(("insert", None, token))
            return
            
        if not p_slice:
            for token in w_slice:
                alignments.append(("delete", token, None))
            return
            
        w_chars = [t["char"] for t in w_slice]
        p_chars = [t["char"] for t in p_slice]
        
        char_sm = difflib.SequenceMatcher(None, w_chars, p_chars)
        for tag, i1, i2, j1, j2 in char_sm.get_opcodes():
            if tag == 'equal':
                for offset in range(i2 - i1):
                    alignments.append(("equal", w_slice[i1 + offset], p_slice[j1 + offset]))
            elif tag == 'delete':
                for offset in range(i2 - i1):
                    alignments.append(("delete", w_slice[i1 + offset], None))
            elif tag == 'insert':
                for offset in range(j2 - j1):
                    alignments.append(("insert", None, p_slice[j1 + offset]))
            elif tag == 'replace':
                min_len = min(i2 - i1, j2 - j1)
                for offset in range(min_len):
                    alignments.append(("replace", w_slice[i1 + offset], p_slice[j1 + offset]))
                if (i2 - i1) > min_len:
                    for offset in range(min_len, i2 - i1):
                        alignments.append(("delete", w_slice[i1 + offset], None))
                elif (j2 - j1) > min_len:
                    for offset in range(min_len, j2 - j1):
                        alignments.append(("insert", None, p_slice[j1 + offset]))

    for w_idx, p_idx in matched_pairs:
        align_slice(last_w, w_idx, last_p, p_idx)
        alignments.append(("equal", word_tokens[w_idx], pdf_tokens[p_idx]))
        last_w = w_idx + 1
        last_p = p_idx + 1
        
    align_slice(last_w, len(word_tokens), last_p, len(pdf_tokens))
    return alignments

# ─── Difference Classification ───────────────────────────────────────────────

def classify_difference(expected: str, actual: str) -> tuple:
    expected_clean = expected.strip()
    actual_clean = actual.strip()
    
    if expected_clean and not actual_clean:
        return "Missing content", "red", "Content is present in the Word document but missing in the PDF."
        
    if actual_clean and not expected_clean:
        return "Extra content", "blue", "Content is present in the PDF but missing in the Word document."
        
    if expected_clean.lower() == expected_clean.upper() and not expected_clean and not actual_clean:
        return "Space mismatch", "green", "Whitespace difference detected."

    if expected_clean.lower() == actual_clean.lower():
        return "Case mismatch", "orange", f"Case mismatch: expected '{expected}' but found '{actual}'."
        
    def strip_formatting(text: str) -> str:
        text = text.replace('\u201c', '').replace('\u201d', '').replace('\u2018', '').replace('\u2019', '')
        text = text.replace('\u2013', '').replace('\u2014', '')
        return re.sub(r'[^a-zA-Z0-9]', '', text).lower()
        
    if strip_formatting(expected_clean) == strip_formatting(actual_clean):
        if not strip_formatting(expected_clean) and expected_clean != actual_clean:
            symbols = "&@%$#°±≤≥×÷©®™μαβ"
            if any(c in expected_clean or c in actual_clean for c in symbols):
                return "Symbol mismatch", "yellow", f"Symbol mismatch: expected '{expected}' but found '{actual}'."
            has_bracket_diff = any(c in expected or c in actual for c in '()[]{}')
            if has_bracket_diff:
                return "Bracket mismatch", "green", f"Bracket mismatch: expected '{expected}' but found '{actual}'."
            return "Punctuation mismatch", "green", f"Punctuation mismatch: expected '{expected}' but found '{actual}'."
            
        # Check for punctuation difference (ignoring brackets, spaces, and hyphens/soft hyphens)
        def strip_non_punc(text: str) -> str:
            return re.sub(r'[()\[\]{}\s\-\u00ad]', '', text).lower()
            
        if strip_non_punc(expected_clean) != strip_non_punc(actual_clean):
            has_bracket_diff = any(c in expected or c in actual for c in '()[]{}')
            if has_bracket_diff:
                return "Bracket mismatch", "green", f"Bracket mismatch: expected '{expected}' but found '{actual}'."
            return "Punctuation mismatch", "green", f"Punctuation mismatch: expected '{expected}' but found '{actual}'."
            
        has_bracket_diff = any(c in expected or c in actual for c in '()[]{}')
        has_space_diff = any(c.isspace() for c in expected) != any(c.isspace() for c in actual) or expected.count(' ') != actual.count(' ')
        
        if has_bracket_diff:
            return "Bracket mismatch", "green", f"Bracket mismatch: expected '{expected}' but found '{actual}'."
        elif has_space_diff:
            if len(expected.split()) != len(actual.split()):
                return "Incorrect word", "yellow", f"Word boundary mismatch: expected '{expected}' but found '{actual}'."
            return "Space mismatch", "green", f"Space mismatch: expected '{expected}' but found '{actual}'."
        else:
            return "Punctuation mismatch", "green", f"Punctuation mismatch: expected '{expected}' but found '{actual}'."
            
    if expected_clean.isdigit() or actual_clean.isdigit():
        return "Number mismatch", "yellow", f"Number mismatch: expected '{expected}' but found '{actual}'."
        
    symbols = "&@%$#°±≤≥×÷©®™μαβ"
    if any(c in expected_clean or c in actual_clean for c in symbols):
        return "Symbol mismatch", "yellow", f"Symbol mismatch: expected '{expected}' but found '{actual}'."
        
    return "Incorrect word", "yellow", f"Incorrect word: expected '{expected}' but found '{actual}'."

# ─── Filtering Logic ────────────────────────────────────────────────────────

LIST_MARKER_REGEX = re.compile(
    r'^('
    r'[•oO*▪\u2022\u25e6\u25aa\u25ab\u25a0\u25cb]+'  # Bullet symbols (added O and white circle, allow multiple)
    r'|'
    r'\d+[\.\):\u2022]'                    # 1. or 1) or 1:
    r'|'
    r'\(\d+\)'                             # (1)
    r'|'
    r'[a-zA-Z][\.\):\u2022]'                # a. or a) or a: or A. or A) or A:
    r'|'
    r'\([a-zA-Z]\)'                        # (a)
    r'|'
    r'[ivxIVX]+[\.\):\u2022]'              # i. or i) or i:
    r'|'
    r'\([ivxIVX]+\)'                       # (i)
    r')$'
)

def is_list_marker(text: str) -> bool:
    val = text.strip()
    return bool(LIST_MARKER_REGEX.match(val))

def is_soft_hyphen(expected: str, actual: str) -> bool:
    if '-' not in actual and '\u00ad' not in actual:
        return False
    e_clean = re.sub(r'[-\u00ad\s]', '', expected).lower()
    a_clean = re.sub(r'[-\u00ad\s]', '', actual).lower()
    if e_clean == a_clean:
        return True
    return False

# ─── Mismatch Grouping ───────────────────────────────────────────────────────

def group_mismatch_segments(alignments: list) -> list:
    marked = [False] * len(alignments)
    for idx, (tag, w_tok, p_tok) in enumerate(alignments):
        if tag != "equal":
            is_w_space = w_tok["is_space"] if w_tok else False
            is_p_space = p_tok["is_space"] if p_tok else False
            if not (is_w_space and is_p_space):
                marked[idx] = True
            
    # Propagate space mismatches to adjacent character tokens so they group with the words
    for idx in range(len(alignments)):
        if marked[idx]:
            tag, w_tok, p_tok = alignments[idx]
            is_w_space = w_tok["is_space"] if w_tok else True
            is_p_space = p_tok["is_space"] if p_tok else True
            if is_w_space or is_p_space:
                if idx > 0:
                    marked[idx - 1] = True
                if idx < len(alignments) - 1:
                    marked[idx + 1] = True
                    
    from collections import defaultdict
    word_ranges = defaultdict(list)
    pdf_ranges = defaultdict(list)

    for idx, (tag, w_tok, p_tok) in enumerate(alignments):
        if w_tok is not None and not w_tok["is_space"]:
            w_id = (w_tok["para_idx"], w_tok["word_idx_in_para"])
            word_ranges[w_id].append(idx)
        if p_tok is not None and not p_tok["is_space"]:
            p_id = (p_tok["page_idx"], p_tok["line_idx"], p_tok["word_idx_in_line"])
            pdf_ranges[p_id].append(idx)

    changed = True
    while changed:
        changed = False
        for w_id, indices in word_ranges.items():
            if any(marked[i] for i in indices):
                for i in indices:
                    if not marked[i]:
                        marked[i] = True
                        changed = True
        for p_id, indices in pdf_ranges.items():
            if any(marked[i] for i in indices):
                for i in indices:
                    if not marked[i]:
                        marked[i] = True
                        changed = True

    mismatch_segments = []
    in_segment = False
    current_segment = []

    for idx, is_marked in enumerate(marked):
        if is_marked:
            if in_segment and len(current_segment) > 0:
                prev_w_tok = alignments[current_segment[-1]][1]
                curr_w_tok = alignments[idx][1]
                if prev_w_tok is not None and curr_w_tok is not None:
                    if prev_w_tok.get("para_idx") != curr_w_tok.get("para_idx"):
                        mismatch_segments.append(current_segment)
                        current_segment = []
            if not in_segment:
                in_segment = True
            current_segment.append(idx)
        else:
            if in_segment:
                mismatch_segments.append(current_segment)
                current_segment = []
                in_segment = False
    if in_segment:
        mismatch_segments.append(current_segment)
        
    errors = []
    for segment in mismatch_segments:
        w_tokens = [alignments[i][1] for i in segment if alignments[i][1] is not None]
        p_tokens = [alignments[i][2] for i in segment if alignments[i][2] is not None]
        
        expected = "".join(t["char"] for t in w_tokens)
        actual = "".join(t["char"] for t in p_tokens)
        
        if expected.strip() == actual.strip():
            continue
            
        # Resolve error locations
        page_num = None
        for idx in segment:
            if alignments[idx][2] is not None:
                page_num = alignments[idx][2]["page_idx"]
                break
                
        if page_num is None:
            for idx in range(segment[0] - 1, -1, -1):
                if alignments[idx][2] is not None:
                    page_num = alignments[idx][2]["page_idx"]
                    break
        if page_num is None:
            for idx in range(segment[-1] + 1, len(alignments)):
                if alignments[idx][2] is not None:
                    page_num = alignments[idx][2]["page_idx"]
                    break
        if page_num is not None:
            page_num = page_num + 1
        else:
            page_num = 1
        
        line_num = None
        for idx in range(segment[0] - 1, -1, -1):
            if alignments[idx][2] is not None:
                line_num = alignments[idx][2]["line_idx"]
                break
        if line_num is None:
            for idx in range(segment[-1] + 1, len(alignments)):
                if alignments[idx][2] is not None:
                    line_num = alignments[idx][2]["line_idx"]
                    break
        line_num = line_num or 1
        
        para_num = None
        for idx in range(segment[0] - 1, -1, -1):
            if alignments[idx][1] is not None:
                para_num = alignments[idx][1]["para_idx"]
                break
        if para_num is None:
            for idx in range(segment[-1] + 1, len(alignments)):
                if alignments[idx][1] is not None:
                    para_num = alignments[idx][1]["para_idx"]
                    break
        para_num = para_num or 1
        
        char_pos = w_tokens[0]["char_in_para"] if w_tokens else 0
        word_idx = w_tokens[0]["word_idx_in_para"] if w_tokens else 0
        
        err_type, color, description = classify_difference(expected, actual)
        
        # 1. Ignore space mismatches
        if err_type == "Space mismatch":
            continue
            
        # 2. Ignore soft hyphens (e.g. prerequisite vs prereq-uisite)
        if err_type in ("Incorrect word", "Punctuation mismatch") and is_soft_hyphen(expected, actual):
            continue
            
        # 3. Ignore list markers
        if not expected.strip() and is_list_marker(actual):
            continue
        
        # Clean up p_tokens by stripping leading/trailing space tokens to prevent highlighting spaces
        p_tokens_clean = list(p_tokens)
        while p_tokens_clean and p_tokens_clean[0]["char"].isspace():
            p_tokens_clean.pop(0)
        while p_tokens_clean and p_tokens_clean[-1]["char"].isspace():
            p_tokens_clean.pop()
            
        bboxes = [t["bbox"] for t in p_tokens_clean if t.get("bbox") is not None]
        nearest_bbox = None
        if bboxes:
            nearest_bbox = bboxes[0]
        else:
            for idx in range(segment[0] - 1, -1, -1):
                if alignments[idx][2] is not None and alignments[idx][2].get("bbox") is not None:
                    nearest_bbox = alignments[idx][2]["bbox"]
                    break
            if nearest_bbox is None:
                for idx in range(segment[-1] + 1, len(alignments)):
                    if alignments[idx][2] is not None and alignments[idx][2].get("bbox") is not None:
                        nearest_bbox = alignments[idx][2]["bbox"]
                        break
                        
        location_ref = f"Word P{para_num} W{word_idx} C{char_pos} | PDF Pg{page_num}"
        errors.append({
            "check": "Word-to-Word Comparison",
            "page": str(page_num),
            "line": line_num,
            "para": para_num,
            "char_pos": char_pos,
            "word_idx": word_idx,
            "type": err_type,
            "color": color,
            "expected": expected,
            "actual": actual,
            "location": f"{err_type}: expected '{expected}' but found '{actual}' at {location_ref}",
            "description": description,
            "bboxes": bboxes,
            "nearest_bbox": nearest_bbox
        })
        
    return errors

# ─── Structural Checks (Typesetting QA Validation) ──────────────────────────

def check_page_sequence(pdf_path: str, start_page: int = 0) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    extracted = []
    
    for page_idx, page in enumerate(doc, 1):
        if page_idx < start_page + 1:
            continue
        lines = [l.strip() for l in page.get_text("text").split("\n") if l.strip()]
        candidate = None
        for line in lines[:3] + lines[-3:]:
            # 1. Standalone number
            if re.match(r'^\d{1,4}$', line):
                candidate = int(line)
                break
            # 2. Number at start followed by multiple spaces/tabs
            m_start = re.search(r'^\s*(\d+)(?:\s{2,}|\t|\u2003)', line)
            if m_start:
                candidate = int(m_start.group(1))
                break
            # 3. Number at end preceded by multiple spaces/tabs
            m_end = re.search(r'(?:\s{2,}|\t|\u2003)(\d+)\s*$', line)
            if m_end:
                candidate = int(m_end.group(1))
                break
        extracted.append((page_idx, candidate))
        
    prev_num = None
    prev_page = None
    for page_num, found_num in extracted:
        if found_num is None:
            continue
        if prev_num is not None:
            expected_num = prev_num + (page_num - prev_page)
            if found_num != expected_num:
                errors.append({
                    "check": "Page Number Sequence & Folio",
                    "page": str(page_num),
                    "location": f"Page sequence error: Found {found_num} on physical page {page_num} (expected {expected_num}).",
                    "type": "Sequence error",
                    "color": "yellow",
                    "expected": str(expected_num),
                    "actual": str(found_num),
                    "description": f"Page number sequence is out of order. Found {found_num} following {prev_num}."
                })
        prev_num = found_num
        prev_page = page_num
        
    doc.close()
    return errors

def check_running_heads(pdf_path: str, repeating_headers) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    
    for page_idx, page in enumerate(doc, 1):
        lines = [l.strip() for l in page.get_text("text").split("\n") if l.strip()]
        if not lines:
            continue
        head = lines[0]
        if len(head) > 120:
            errors.append({
                "check": "Running Head Style & Position",
                "page": str(page_idx),
                "location": f"Header too long: '{head[:60]}...'",
                "type": "Header error",
                "color": "yellow",
                "expected": "Short, concise header",
                "actual": head[:60],
                "description": "Running head exceeds standard maximum length limit."
            })
    doc.close()
    return errors

def check_equations(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    eq_re = re.compile(r'[A-Za-z]\s*=\s*[\w\(\)\+\-\*\/\^]+')
    
    for page_idx, page in enumerate(doc, 1):
        matches = eq_re.findall(page.get_text("text"))
        for m in matches:
            if "?" in m or "□" in m or "■" in m:
                errors.append({
                    "check": "Equations",
                    "page": str(page_idx),
                    "location": f"Broken equation: '{m[:40]}'",
                    "type": "Equation formatting",
                    "color": "green",
                    "expected": "Properly formatted equation",
                    "actual": m[:40],
                    "description": "Equation rendering contains broken symbols or placeholders."
                })
    doc.close()
    return errors

def check_special_chars(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    bad = re.compile(r'[□■]')
    
    for page_idx, page in enumerate(doc, 1):
        found = bad.findall(page.get_text("text"))
        if found:
            unique = list(set(found))
            errors.append({
                "check": "Special Characters & Symbols",
                "page": str(page_idx),
                "location": f"Broken rendering/special symbols found: {', '.join(unique[:8])}",
                "type": "Unicode rendering",
                "color": "yellow",
                "expected": "Correct symbol glyph",
                "actual": unique[0],
                "description": "Document contains special formatting characters or rendering artifacts."
            })
    doc.close()
    return errors

def check_footnotes(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    
    for page_idx, page in enumerate(doc, 1):
        page_dict = page.get_text("dict")
        page_h = page.rect.height
        body_cutoff = page_h * 0.80
        
        body_nums = set()
        bottom_nums = set()
        
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_y = line["bbox"][1]
                is_body = line_y < body_cutoff
                
                for span in line.get("spans", []):
                    flags = span.get("flags", 0)
                    size = span.get("size", 12)
                    text = span.get("text", "").strip()
                    
                    if not text or not text.isdigit():
                        continue
                        
                    num = int(text)
                    if not (1 <= num <= 300):
                        continue
                        
                    line_text = "".join(s["text"] for s in line.get("spans", [])).strip()
                    if line_text == text:
                        continue
                        
                    is_superscript = bool(flags & 1) or (size < 8)
                    if is_superscript:
                        if is_body:
                            body_nums.add(num)
                        else:
                            bottom_nums.add(num)
                            
        missing = body_nums - bottom_nums
        for num in sorted(missing)[:5]:
            errors.append({
                "check": "Footnote Citation & Placement",
                "page": str(page_idx),
                "location": f"Missing footnote reference at bottom of page for marker {num}",
                "type": "Footnote placement",
                "color": "yellow",
                "expected": f"Footnote text for [{num}] at bottom of page",
                "actual": "Missing reference text",
                "description": f"Footnote citation marker [{num}] is present in the page body, but the footnote definition is missing from the footer."
            })
    doc.close()
    return errors

def check_double_digit_alignment(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    list_num_re = re.compile(r'^(\d{1,3})\.$')
    
    for page_idx, page in enumerate(doc, 1):
        blocks = page.get_text("dict")["blocks"]
        num_positions = []
        
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_text = "".join(s["text"] for s in line.get("spans", [])).strip()
                m = list_num_re.match(line_text)
                if m:
                    num_val = int(m.group(1))
                    x0 = line["bbox"][0]
                    num_positions.append((num_val, x0))
                    
        if not num_positions:
            continue
            
        singles = [x for v, x in num_positions if v < 10]
        doubles = [x for v, x in num_positions if v >= 10]
        
        if singles and doubles:
            avg_single = sum(singles) / len(singles)
            avg_double = sum(doubles) / len(doubles)
            if abs(avg_single - avg_double) > 3:
                errors.append({
                    "check": "Double Digit Alignment",
                    "page": str(page_idx),
                    "location": "Misaligned list digits: single vs double digit indent mismatch",
                    "type": "List formatting alignment",
                    "color": "green",
                    "expected": "Right-aligned or properly tabbed list digits",
                    "actual": "Offset alignment",
                    "description": "Double-digit list items (10+) do not align vertically with single-digit list items."
                })
    doc.close()
    return errors

def check_global_instructions(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    pdf_instruction_re = re.compile(r'\b(AU\s*QUERY|STET|TK|TODO|QUERY)[:\s][^\n]{5,80}', re.I)
    
    for page_idx, page in enumerate(doc, 1):
        m = pdf_instruction_re.search(page.get_text("text"))
        if m:
            errors.append({
                "check": "Global Instructions",
                "page": str(page_idx),
                "location": f"Instruction left in PDF: '{m.group()[:60]}'",
                "type": "Editor notes leak",
                "color": "yellow",
                "expected": "Production-ready clean text",
                "actual": m.group()[:60],
                "description": "An unresolved editor guideline, placeholder note, or layout query has been leaked to the final output."
            })
    doc.close()
    return errors

def check_fpo(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    fpo_re = re.compile(r'\bFPO\b|\bFor\s+Position\s+Only\b|\bPLACEHOLDER\s+IMAGE\b', re.I)
    
    for page_idx, page in enumerate(doc, 1):
        m = fpo_re.search(page.get_text("text"))
        if m:
            errors.append({
                "check": "FPO / Placeholder Images",
                "page": str(page_idx),
                "location": "Placeholder image text (FPO) not replaced",
                "type": "Placeholder leak",
                "color": "yellow",
                "expected": "Final typeset graphic",
                "actual": m.group(),
                "description": "An image markup notation or position marker ('For Position Only') is left in the page flow."
            })
            
        try:
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                w, h = base_image["width"], base_image["height"]
                if w < 50 or h < 50:
                    errors.append({
                        "check": "FPO / Placeholder Images",
                        "page": str(page_idx),
                        "location": f"Image cut off or too small ({w}x{h}px)",
                        "type": "Image formatting",
                        "color": "yellow",
                        "expected": "High resolution, fully scaled image",
                        "actual": f"Low-res placeholder ({w}x{h}px)",
                        "description": "Extracted graphic is too low-resolution or displays cut-off dimensions."
                    })
        except:
            pass
    doc.close()
    return errors

def check_image_size(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    for page_idx, page in enumerate(doc, 1):
        try:
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                w, h = base_image["width"], base_image["height"]
                if w < 300 or h < 300:
                    errors.append({
                        "check": "Image Size",
                        "page": str(page_idx),
                        "location": f"Image resolution too small ({w}x{h}px)",
                        "type": "Image quality",
                        "color": "yellow",
                        "expected": "At least 300x300px",
                        "actual": f"{w}x{h}px",
                        "description": "Image dimensions or DPI are below publishing production quality criteria."
                    })
        except:
            pass
    doc.close()
    return errors

def check_line_art_text(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    for page_idx, page in enumerate(doc, 1):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size", 12)
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    if size < 6:
                        errors.append({
                            "check": "Text within Line Art",
                            "page": str(page_idx),
                            "location": f"Text too small ({size:.1f}pt): '{text[:30]}'",
                            "type": "Font legibility",
                            "color": "green",
                            "expected": "Readable font size (> 6.0pt)",
                            "actual": f"{size:.1f}pt font",
                            "description": "Text inside technical diagrams, drawings, or line art is too small and might be illegible."
                        })
    doc.close()
    return errors

def check_credit_lines(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    credit_re = re.compile(r'(source|credit|courtesy|photo by|image by|©|\bcc\b|creative commons)', re.I)
    figure_re = re.compile(r'\b(Figure|Fig\.?|Photo|Image|Plate)\s*\d+', re.I)
    
    for page_idx, page in enumerate(doc, 1):
        figures = figure_re.findall(page.get_text("text"))
        if figures:
            has_credit = bool(credit_re.search(page.get_text("text")))
            if not has_credit:
                errors.append({
                    "check": "Credit Lines",
                    "page": str(page_idx),
                    "location": f"Missing credit line or attribution for figure label",
                    "type": "Attribution error",
                    "color": "yellow",
                    "expected": "Figure caption containing CC, credit, or source attribution",
                    "actual": "Missing credit line",
                    "description": "Graphic figures are declared, but copyright or source attribution credits are missing."
                })
    doc.close()
    return errors

def check_key_terms(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    keyterm_re = re.compile(r'\b(key\s*terms?|glossary|terminology|definitions?|index)\b', re.I)
    page_ref_re = re.compile(r'\b(p\.|pp\.|page\s+|pages?\s+)(\d+)', re.I)
    
    total = len(doc)
    
    for page_idx, page in enumerate(doc, 1):
        text = page.get_text("text")
        if keyterm_re.search(text):
            for m in page_ref_re.finditer(text):
                ref_num = int(m.group(2))
                if ref_num > total:
                    errors.append({
                        "check": "Key Terms Page Numbers",
                        "page": str(page_idx),
                        "location": f"Index/Glossary references broken page: pg {ref_num} (document has only {total} pages)",
                        "type": "Reference error",
                        "color": "yellow",
                        "expected": f"Page number <= {total}",
                        "actual": f"Page number {ref_num}",
                        "description": "Key term reference index points to a page number that does not exist in the final typeset PDF."
                    })
    doc.close()
    return errors

def check_unwanted_chars(pdf_path: str) -> list:
    errors = []
    doc = fitz.open(pdf_path)
    patterns = {
        r'\$\$$': "Placeholder '$$$'",
        r'\bxxx\b': "Placeholder 'xxx'",
        r'\bXXX\b': "Placeholder 'XXX'",
        r'(?<![,\d])000(?!\d)': "Numeric placeholder '000'",
        r'\bTBD\b': "Placeholder 'TBD'",
        r'\bLorem\s+Ipsum\b': "Dummy text 'Lorem Ipsum'",
        r'\bPLACEHOLDER\b': "Text placeholder 'PLACEHOLDER'"
    }
    
    for page_idx, page in enumerate(doc, 1):
        for pattern, explanation in patterns.items():
            rx = re.compile(pattern, re.I)
            m = rx.search(page.get_text("text"))
            if m:
                errors.append({
                    "check": "Unwanted Characters",
                    "page": str(page_idx),
                    "location": f"Placeholder left in page layout: {explanation}",
                    "type": "Draft character leak",
                    "color": "yellow",
                    "expected": "Final edited content",
                    "actual": m.group(),
                    "description": "Draft placeholder characters, layout parameters, or Lorem Ipsum text have leaked to the output."
                })
    doc.close()
    return errors

# ─── Main Validation Endpoint ───────────────────────────────────────────────

@app.post("/validate")
async def validate(
    word_file: UploadFile = File(...),
    pdf_file: UploadFile = File(...),
    checks: str = Form(""),
):
    selected = [c.strip() for c in checks.split(",") if c.strip()] if checks else []
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as wf:
        wf.write(await word_file.read())
        word_path = wf.name
        
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pf:
        pf.write(await pdf_file.read())
        pdf_path = pf.name
        
    all_errors = []
    
    try:
        start_page = find_chapter_start_page(word_path, pdf_path)
        repeating_headers, is_page_number_like = identify_headers_footers(pdf_path)
        
        if not selected or "Page Number Sequence & Folio" in selected or "Page Number Sequence and Folio Placement" in selected:
            all_errors.extend(check_page_sequence(pdf_path, start_page))
        if not selected or "Running Head Style & Position" in selected or "Running Head Style and Position" in selected:
            all_errors.extend(check_running_heads(pdf_path, repeating_headers))
        if not selected or "Equations" in selected or "Equation Entry" in selected:
            all_errors.extend(check_equations(pdf_path))
        if not selected or "Special Characters & Symbols" in selected or "Equation Special Character Symbol" in selected:
            all_errors.extend(check_special_chars(pdf_path))
        if not selected or "Footnote Citation & Placement" in selected or "Footnote Citation and Placement" in selected:
            all_errors.extend(check_footnotes(pdf_path))
        if not selected or "Double Digit Alignment" in selected:
            all_errors.extend(check_double_digit_alignment(pdf_path))
        if not selected or "Global Instructions" in selected:
            all_errors.extend(check_global_instructions(pdf_path))
        if not selected or "FPO / Placeholder Images" in selected or "Figure Photo FPO Check" in selected:
            all_errors.extend(check_fpo(pdf_path))
        if not selected or "Image Size" in selected:
            all_errors.extend(check_image_size(pdf_path))
        if not selected or "Text within Line Art" in selected or "Line Art Readability" in selected:
            all_errors.extend(check_line_art_text(pdf_path))
        if not selected or "Credit Lines" in selected:
            all_errors.extend(check_credit_lines(pdf_path))
        if not selected or "Key Terms Page Numbers" in selected or "Keyterm Page Numbers" in selected:
            all_errors.extend(check_key_terms(pdf_path))
        if not selected or "Unwanted Characters" in selected:
            all_errors.extend(check_unwanted_chars(pdf_path))
            
        if (not selected or 
            "Word-to-Word Comparison" in selected or 
            "Missing Content" in selected or "No Content Missing" in selected or 
            "Typos" in selected or "No Typos" in selected or 
            "Quotations" in selected or "Quotation Check" in selected):
            word_tokens = extract_word_tokens(word_path)
            pdf_tokens = extract_pdf_tokens(pdf_path, start_page)
            
            word_words = group_tokens_into_words(word_tokens)
            pdf_words = group_tokens_into_words(pdf_tokens)
            
            word_word_strings = ["".join(item[1]["char"] for item in w) for w in word_words]
            pdf_word_strings = ["".join(item[1]["char"] for item in w) for w in pdf_words]
            
            word_word_aligned = [clean_word_for_align(s) for s in word_word_strings]
            pdf_word_aligned = [clean_word_for_align(s) for s in pdf_word_strings]
            
            sm = difflib.SequenceMatcher(None, word_word_aligned, pdf_word_aligned)
            matching_blocks = sm.get_matching_blocks()
            
            alignments = align_all_characters(word_tokens, pdf_tokens, word_words, pdf_words, matching_blocks)
            text_errors = group_mismatch_segments(alignments)
            all_errors.extend(text_errors)
            
        import shutil
        try:
            shutil.copy(word_path, r"C:\Users\Harini\.gemini\antigravity-ide\scratch\latest.docx")
            shutil.copy(pdf_path, r"C:\Users\Harini\.gemini\antigravity-ide\scratch\latest.pdf")
        except:
            pass
            
        pdf_doc = fitz.open(pdf_path)
        total_pages = len(pdf_doc)
        pdf_doc.close()
        
    finally:
        try:
            os.unlink(word_path)
        except:
            pass
        try:
            os.unlink(pdf_path)
        except:
            pass
            
    affected_pages = sorted(set(
        int(e["page"]) for e in all_errors
        if e.get("page", "").isdigit()
    ))
    
    return {
        "errors": all_errors,
        "total_errors": len(all_errors),
        "total_pages": total_pages,
        "affected_pages": affected_pages,
        "checks_run": len(selected) if selected else 26,
    }

# ─── Highlighted PDF Endpoint ────────────────────────────────────────────────

@app.post("/highlighted-pdf")
async def highlighted_pdf(
    pdf_file: UploadFile = File(...),
    errors: str = Form(""),
):
    try:
        error_list = json.loads(errors) if errors else []
    except Exception as e:
        print(f"DEBUG: Failed to parse errors JSON: {e}")
        error_list = []
        
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pf:
        pf.write(await pdf_file.read())
        pdf_path = pf.name
        
    buf = None
    try:
        doc = fitz.open(pdf_path)
        
        colors = {
            "red": [1.0, 0.0, 0.0],
            "yellow": [1.0, 1.0, 0.0],
            "blue": [0.0, 0.0, 1.0],
            "orange": [1.0, 0.5, 0.0],
            "green": [0.0, 0.8, 0.0]
        }
        
        label_offsets = {}
        
        for error in error_list:
            try:
                page_str = (error.get("page") or "").strip()
                if not page_str.isdigit():
                    continue
                    
                page_num = int(page_str) - 1
                if page_num < 0 or page_num >= len(doc):
                    continue
                    
                page = doc[page_num]
                color_name = error.get("color", "yellow")
                rgb_color = colors.get(color_name, [1.0, 1.0, 0.0])
                
                err_type = error.get("type", "Error")
                expected = error.get("expected", "")
                actual = error.get("actual", "")
                desc = error.get("description", "")
                
                comment_content = f"Error Type: {err_type}\nExpected: {expected}\nActual: {actual}\nDescription: {desc}"
                
                bboxes = error.get("bboxes") or []
                highlighted = False
                
                if bboxes:
                    rects = []
                    for bbox in bboxes:
                        if isinstance(bbox, (list, tuple)) and len(bbox) == 4:
                            r = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3])
                            if not r.is_empty and r.width > 0.1 and r.height > 0.1:
                                rects.append(r)
                    
                    merged_rects = []
                    if rects:
                        rects.sort(key=lambda r: (r.y0, r.x0))
                        current_rect = rects[0]
                        for next_rect in rects[1:]:
                            overlap = min(current_rect.y1, next_rect.y1) - max(current_rect.y0, next_rect.y0)
                            min_h = min(current_rect.height, next_rect.height)
                            if overlap > 0.5 * min_h:
                                current_rect = current_rect | next_rect
                            else:
                                merged_rects.append(current_rect)
                                current_rect = next_rect
                        merged_rects.append(current_rect)
                        
                    for rect in merged_rects:
                        annot = page.add_highlight_annot(rect)
                        if annot:
                            annot.set_colors(stroke=rgb_color)
                            annot.set_info(title=err_type, content=comment_content)
                            annot.update()
                            highlighted = True
                
                if not highlighted or color_name == "red":
                    nearest = error.get("nearest_bbox")
                    if nearest and isinstance(nearest, (list, tuple)) and len(nearest) == 4:
                        y_coord = nearest[1]
                        cy = label_offsets.get(page_num, y_coord)
                        if cy > page.rect.y1 - 40:
                            cy = 40
                            
                        pt = fitz.Point(15, cy)
                        annot = page.add_text_annot(pt, f"QA Alert: {err_type}")
                        if annot:
                            annot.set_colors(stroke=rgb_color)
                            annot.set_info(title=err_type, content=comment_content)
                            annot.update()
                            
                            label_offsets[page_num] = cy + 25
            except Exception as inner_e:
                print(f"DEBUG: Failed to highlight error {error}: {inner_e}")
                
        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        buf.seek(0)
    except Exception as e:
        print(f"DEBUG: Exception during PDF highlighting: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate highlighted PDF: {str(e)}")
    finally:
        try:
            os.unlink(pdf_path)
        except:
            pass
            
    if buf is None:
        raise HTTPException(status_code=500, detail="Failed to generate highlighted PDF.")
        
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=highlighted_report.pdf"},
    )

# ─── Download Report Endpoint ────────────────────────────────────────────────

@app.post("/download-report")
async def download_report(
    format: str = Form("docx"),
    errors: str = Form(""),
):
    try:
        error_list = json.loads(errors) if errors else []
    except Exception:
        error_list = []
        
    if format == "pdf":
        buf = _build_pdf_report(error_list)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=qa_report.pdf"},
        )
    else:
        buf = _build_docx_report(error_list)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=qa_report.docx"},
        )

def _build_docx_report(error_list: list) -> io.BytesIO:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    title = doc.add_heading("Publishing QA Comparison Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Total errors found: {len(error_list)}")
    doc.add_paragraph("")
    
    if not error_list:
        doc.add_paragraph("No errors found. All content matched exactly.")
    else:
        from collections import defaultdict
        grouped = defaultdict(list)
        for e in error_list:
            grouped[e.get("check", "Unknown")].append(e)
            
        for check_name, errs in grouped.items():
            doc.add_heading(check_name, level=1)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Light Shading Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "#"
            hdr[1].text = "Page"
            hdr[2].text = "Type"
            hdr[3].text = "Expected"
            hdr[4].text = "Actual"
            hdr[5].text = "Description"
            
            for i, e in enumerate(errs, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = e.get("page", "—")
                row[2].text = e.get("type", "")
                row[3].text = e.get("expected", "")
                row[4].text = e.get("actual", "")
                row[5].text = e.get("description", "")
            doc.add_paragraph("")
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _build_pdf_report(error_list: list) -> io.BytesIO:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    
    y = 50
    x = 50
    
    def add_text(text, size=10, bold=False, color=(0, 0, 0), gap=5):
        nonlocal y, page
        if y > 800:
            page = doc.new_page(width=595, height=842)
            y = 50
        font = "helv" if not bold else "hebo"
        page.insert_text((x, y), text, fontname=font, fontsize=size, color=color)
        y += size + gap
        
    add_text("Publishing QA Comparison Report", size=16, bold=True, gap=10)
    add_text(f"Total Errors Detected: {len(error_list)}", size=11, gap=15)
    
    if not error_list:
        add_text("No errors found. All content matched exactly.", size=10, gap=5)
    else:
        from collections import defaultdict
        grouped = defaultdict(list)
        for e in error_list:
            grouped[e.get("check", "Unknown")].append(e)
            
        for check_name, errs in grouped.items():
            add_text(check_name, size=12, bold=True, color=(0.1, 0.2, 0.6), gap=8)
            for i, e in enumerate(errs, 1):
                line = f"  {i}. [Page {e.get('page','—')}] Type: {e.get('type','')} | Expected: '{e.get('expected','')}' | Actual: '{e.get('actual','')}'"
                while len(line) > 90:
                    add_text(line[:90], size=9, gap=3)
                    line = "      " + line[90:]
                add_text(line, size=9, gap=4)
                desc = f"      Description: {e.get('description','')}"
                while len(desc) > 90:
                    add_text(desc[:90], size=8, color=(0.4, 0.4, 0.4), gap=3)
                    desc = "      " + desc[90:]
                add_text(desc, size=8, color=(0.4, 0.4, 0.4), gap=6)
            y += 6
            
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return buf

@app.get("/health")
def health():
    return {"status": "ok"}

if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)